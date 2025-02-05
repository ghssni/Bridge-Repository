'''
Script ini digunakan untuk merge QR Bridge Pocket dan Frame BG Tournament
- Pertama QR akan digabungkan dengan Frame PNG
- Kedua QR + Frame akan disusun pada microsoft docs sehingga siap untuk diprint

Ukuran yang digunakan pada script ini:
    - QR Code: 810x860 piksel
    - Frame.png: 1080x1350 piksel

Isikan path folder sebelum menjalankan script ini:
    - qr_folder
    - frame_path
    - output_folder 

Sesuaikan juga letak QR code pada frame pada:
    - x_offset
    - y_offset
'''

import os
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image


# Path ke folder yang berisi foto dan frame
qr_folder = ''
frame_path = ''
qr_merge_folder = ''
output_folder = '' # Path untuk docs yang akan diprint

# Pastikan output folder ada
os.makedirs(qr_merge_folder, exist_ok=True)

# Load frame dari file PNG
frame = Image.open(frame_path)
frame_size = frame.size  # Ukuran frame otomatis

qr_size = (810, 860) 

# Loop semua file QR di folder
for file_name in os.listdir(qr_folder):
    if file_name.endswith((".png", ".jpg", ".jpeg")):
        qr_path = os.path.join(qr_folder, file_name)
        
        # Load QR code
        qr_code = Image.open(qr_path)
        
        # Resize QR sesuai ukuran yang diinginkan
        qr_code = qr_code.resize(qr_size)
        
        # Hitung koordinat posisi QR di tengah latar belakang
        x_offset = (1080 - 810) // 2
        y_offset = 400

        # Copy frame supaya tidak overwrite
        result = frame.copy()
        
        # Tempelkan QR di tengah frame
        result.paste(qr_code, (x_offset, y_offset), qr_code if qr_code.mode == 'RGBA' else None)
        
        # Simpan hasil ke folder output
        output_path = os.path.join(qr_merge_folder, file_name)
        result.save(output_path)

print("Proses pertama selesai! Semua QR code telah ditempel pada frame.")

# Path folder QR code
qr_merge_folder = "/Users/ghtyas/Downloads/QR Hasil"

# Buat dokumen baru
doc = Document()

# Set orientasi halaman ke portrait F4 (21.6 cm x 33 cm)
section = doc.sections[-1]
section.page_height = Inches(13)  
section.page_width = Inches(8.5)  
section.orientation = WD_ORIENT.PORTRAIT
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)

# Ambil semua file gambar dalam folder
qr_paths = [os.path.join(qr_merge_folder, file) for file in os.listdir(qr_merge_folder) if file.endswith(('.png', '.jpg', '.jpeg'))]

# Konfigurasi grid 3x3
max_per_row = 3

# Proses penempatan gambar di Word
table = doc.add_table(rows=0, cols=max_per_row)
for i, img_path in enumerate(qr_paths):
    img = Image.open(img_path)

    # Tambahkan baris baru di tabel
    if i % max_per_row == 0:
        row_cells = table.add_row().cells

    # Sisipkan gambar ke dalam sel
    paragraph = row_cells[i % max_per_row].paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(img_path, width=Inches(2.5))  

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Simpan file Word
doc.save(output_folder)